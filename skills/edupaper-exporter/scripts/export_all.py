#!/usr/bin/env python3
"""
EduPaper Exporter — 跨平台论文导出脚本
支持：macOS / Windows / Linux
导出格式：DOCX（Word）、PDF、HTML
中文字体自动适配各平台

用法：
  python export_all.py <input.md> [--formats docx,pdf,html] [--output-dir ./exports]

依赖（按优先级）：
  1. pandoc + xelatex → 全功能（DOCX + PDF + HTML）
  2. pandoc（无 LaTeX） → DOCX + HTML
  3. python-docx → 仅 DOCX
  4. 无任何工具 → 输出安装指引
"""

import sys
import os
import shutil
import subprocess
import argparse
from pathlib import Path


# ─── 平台与工具检测 ────────────────────────────────────────────────────────

def detect_platform():
    if sys.platform == 'darwin':
        return 'mac'
    elif sys.platform.startswith('win'):
        return 'windows'
    else:
        return 'linux'


def check_tool(name):
    """返回工具路径，不可用返回 None"""
    return shutil.which(name)


def check_python_docx():
    try:
        import docx  # noqa
        return True
    except ImportError:
        return False


def detect_capabilities():
    """检测当前环境支持哪些导出方式"""
    caps = {
        'pandoc': check_tool('pandoc'),
        'xelatex': check_tool('xelatex') or check_tool('xetex'),
        'python_docx': check_python_docx(),
        'platform': detect_platform(),
    }
    return caps


# ─── 中文字体选择 ──────────────────────────────────────────────────────────

FONT_MAP = {
    'mac': {
        'body': ['STSong', 'Songti SC', 'PingFang SC', 'Arial Unicode MS'],
        'heading': ['STHeiti', 'Heiti SC', 'PingFang SC'],
    },
    'windows': {
        'body': ['宋体', 'SimSun', '仿宋', 'FangSong'],
        'heading': ['黑体', 'SimHei', '微软雅黑', 'Microsoft YaHei'],
    },
    'linux': {
        'body': ['Noto Serif CJK SC', 'WenQuanYi Bitmap Song', 'AR PL UMing CN'],
        'heading': ['Noto Sans CJK SC', 'WenQuanYi Zen Hei', 'AR PL UKai CN'],
    },
}


def get_cjk_fonts(platform):
    """返回该平台首选中文字体（正文、标题）"""
    fonts = FONT_MAP.get(platform, FONT_MAP['linux'])
    return fonts['body'][0], fonts['heading'][0]


# ─── 导出函数 ──────────────────────────────────────────────────────────────

def export_docx(input_path: Path, output_path: Path, caps: dict) -> bool:
    """导出为 DOCX"""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if caps['pandoc']:
        # 优先用 pandoc（格式更好）
        ref_doc = input_path.parent.parent / 'assets' / 'reference.docx'
        cmd = ['pandoc', str(input_path), '-o', str(output_path)]
        if ref_doc.exists():
            cmd += ['--reference-doc', str(ref_doc)]
        cmd += ['--wrap=none']
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            print(f'  ✓ DOCX → {output_path}')
            return True
        else:
            print(f'  ✗ pandoc DOCX 失败：{result.stderr.strip()}')

    if caps['python_docx']:
        # 备用：python-docx 简单转换
        return _export_docx_python(input_path, output_path)

    print('  ✗ 无法导出 DOCX（需要 pandoc 或 python-docx）')
    return False


def _export_docx_python(input_path: Path, output_path: Path) -> bool:
    """使用 python-docx 将 Markdown 转为 DOCX（基础格式）"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置页面边距（2.54cm）
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line = line.rstrip('\n')
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('> '):
                # 引用块（课堂实录）
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
            elif line.startswith('|'):
                # 表格行：简单处理，转为普通段落
                doc.add_paragraph(line)
            elif line.strip() == '':
                doc.add_paragraph('')
            else:
                doc.add_paragraph(line)

        doc.save(str(output_path))
        print(f'  ✓ DOCX (python-docx) → {output_path}')
        return True
    except Exception as e:
        print(f'  ✗ python-docx 导出失败：{e}')
        return False


def export_pdf(input_path: Path, output_path: Path, caps: dict) -> bool:
    """导出为 PDF（需要 pandoc + xelatex）"""
    if not caps['pandoc']:
        print('  ✗ PDF 导出需要 pandoc（未安装）')
        return False
    if not caps['xelatex']:
        print('  ✗ PDF 导出需要 XeLaTeX（未安装）')
        _print_latex_install_hint(caps['platform'])
        return False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    body_font, heading_font = get_cjk_fonts(caps['platform'])

    # 查找 LaTeX 模板
    template = input_path.parent.parent / 'assets' / 'edupaper.tex'

    cmd = [
        'pandoc', str(input_path),
        '-o', str(output_path),
        '--pdf-engine=xelatex',
        f'-V', f'CJKmainfont={body_font}',
        f'-V', f'CJKsansfont={heading_font}',
        f'-V', 'geometry:margin=2.54cm',
        f'-V', 'fontsize=12pt',
        f'-V', 'linestretch=1.5',
    ]
    if template.exists():
        cmd += ['--template', str(template)]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        print(f'  ✓ PDF → {output_path}')
        return True
    else:
        print(f'  ✗ PDF 导出失败：{result.stderr.strip()[:200]}')
        return False


def export_html(input_path: Path, output_path: Path, caps: dict) -> bool:
    """导出为 HTML（需要 pandoc）"""
    if not caps['pandoc']:
        print('  ✗ HTML 导出需要 pandoc（未安装）')
        return False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    css = input_path.parent.parent / 'assets' / 'style.css'

    cmd = [
        'pandoc', str(input_path),
        '-o', str(output_path),
        '--standalone',
        '--embed-resources',
        '--metadata', f'title={input_path.stem}',
    ]
    if css.exists():
        cmd += ['--css', str(css)]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        print(f'  ✓ HTML → {output_path}')
        return True
    else:
        print(f'  ✗ HTML 导出失败：{result.stderr.strip()[:200]}')
        return False


# ─── 安装提示 ──────────────────────────────────────────────────────────────

def _print_latex_install_hint(platform):
    hints = {
        'mac': 'brew install --cask basictex  # 然后：sudo tlmgr install ctex',
        'windows': '下载 MiKTeX：https://miktex.org/download',
        'linux': 'sudo apt install texlive-xetex texlive-lang-chinese',
    }
    print(f'  💡 安装 LaTeX：{hints.get(platform, hints["linux"])}')


def print_install_guide(caps):
    platform = caps['platform']
    print('\n📦 工具安装指引：')
    if not caps['pandoc']:
        hints = {
            'mac': 'brew install pandoc',
            'windows': 'winget install JohnMacFarlane.Pandoc',
            'linux': 'sudo apt install pandoc',
        }
        print(f'  安装 pandoc：{hints.get(platform, hints["linux"])}')
    if not caps['python_docx']:
        print('  安装 python-docx：pip install python-docx')


# ─── 主流程 ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='EduPaper 跨平台论文导出工具')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('--formats', default='docx',
                        help='导出格式，逗号分隔，如 docx,pdf,html（默认 docx）')
    parser.add_argument('--output-dir', default=None,
                        help='输出目录（默认为输入文件所在目录的 ../../exports/）')
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f'错误：输入文件不存在：{input_path}')
        sys.exit(1)

    # 输出目录
    if args.output_dir:
        output_dir = Path(args.output_dir).resolve()
    else:
        # 默认：.edupaper/exports/
        output_dir = input_path.parent.parent.parent / 'exports'
    output_dir.mkdir(parents=True, exist_ok=True)

    stem = input_path.stem  # 文件名（无扩展名）
    formats = [f.strip().lower() for f in args.formats.split(',')]

    # 检测环境
    caps = detect_capabilities()
    print(f'🔍 平台：{caps["platform"]}')
    print(f'   pandoc：{"✓ " + caps["pandoc"] if caps["pandoc"] else "✗ 未安装"}')
    print(f'   xelatex：{"✓" if caps["xelatex"] else "✗ 未安装（PDF 需要）"}')
    print(f'   python-docx：{"✓" if caps["python_docx"] else "✗ 未安装"}')
    print()

    # 执行导出
    results = {}
    if 'docx' in formats:
        results['docx'] = export_docx(input_path, output_dir / f'{stem}.docx', caps)
    if 'pdf' in formats:
        results['pdf'] = export_pdf(input_path, output_dir / f'{stem}.pdf', caps)
    if 'html' in formats:
        results['html'] = export_html(input_path, output_dir / f'{stem}.html', caps)

    # 汇总
    print()
    success = [fmt for fmt, ok in results.items() if ok]
    failed = [fmt for fmt, ok in results.items() if not ok]
    if success:
        print(f'✅ 导出成功：{", ".join(success).upper()}')
    if failed:
        print(f'⚠️  导出失败：{", ".join(failed).upper()}')
        print_install_guide(caps)

    sys.exit(0 if not failed else 1)


if __name__ == '__main__':
    main()
