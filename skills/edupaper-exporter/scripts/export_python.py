#!/usr/bin/env python3
"""
export_python.py — Python exporter for EduPaper papers.
Produces properly formatted DOCX and HTML with correct Chinese academic
paper formatting: first-line indent, 1.5x line spacing, SimSun/SimHei fonts,
correct heading hierarchy and font sizes.

Usage:
    python3 export_python.py <input.md> <output_dir> [formats]
    formats: all,docx,html (default: all)

Part of edupaper-exporter skill.
"""

import sys
import os
import re
from pathlib import Path

# python-docx imports (available after pip install python-docx)
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
except ImportError:
    Document = None

# ============================================================
# DOCX EXPORT
# ============================================================

def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=12, bold=False):
    """Set font on a run — both Chinese (eastAsia) and Latin."""
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), cn_font)


def set_first_line_indent_chars(paragraph, chars=2):
    """Set first-line indent in character units (Word's '2字符')."""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    # Also set firstLine in twips as fallback (1 char ≈ font size in twips/2)
    # 12pt = 240 twips, 2 chars = 480 twips
    ind.set(qn('w:firstLine'), str(chars * 240))


def set_hanging_indent_chars(paragraph, chars=2):
    """Set hanging indent for references (悬挂缩进)."""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:hangingChars'), str(chars * 100))
    ind.set(qn('w:leftChars'), str(chars * 100))
    ind.set(qn('w:hanging'), str(chars * 240))
    ind.set(qn('w:left'), str(chars * 240))


def md_to_docx(md_path: str, docx_path: str):
    """Convert Markdown to properly formatted DOCX."""
    if Document is None:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # --- Page setup: A4, 上下2.54cm, 左右3.17cm ---
    for section in doc.sections:
        section.page_width = Cm(21.0)    # A4 width
        section.page_height = Cm(29.7)   # A4 height
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)

    # --- Normal style: 宋体小四(12pt), 1.5倍行距 ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # --- Heading styles ---
    # H1 (论文标题): 黑体二号(22pt), 居中
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    rpr1 = h1.element.get_or_add_rPr()
    rf1 = rpr1.get_or_add_rFonts()
    rf1.set(qn('w:eastAsia'), '黑体')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)

    # H2 (一级标题 一、): 黑体三号(16pt), 左对齐, 段前段后0.5行
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    rpr2 = h2.element.get_or_add_rPr()
    rf2 = rpr2.get_or_add_rFonts()
    rf2.set(qn('w:eastAsia'), '黑体')
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)

    # H3 (二级标题 （一）): 黑体四号(14pt), 左对齐
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    rpr3 = h3.element.get_or_add_rPr()
    rf3 = rpr3.get_or_add_rFonts()
    rf3.set(qn('w:eastAsia'), '黑体')
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(0)

    # H4 (三级标题 1.): 宋体小四(12pt)加粗
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Times New Roman'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    rpr4 = h4.element.get_or_add_rPr()
    rf4 = rpr4.get_or_add_rFonts()
    rf4.set(qn('w:eastAsia'), '宋体')
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.line_spacing = 1.5

    # --- Read markdown ---
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    in_references = False  # Track if we're in 参考文献 section

    def add_body_paragraph(text, indent=True):
        """Add a body paragraph with proper formatting."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if indent:
            set_first_line_indent_chars(p, 2)
        # Parse inline bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_run_font(run, '宋体', 'Times New Roman', 12, bold=True)
            else:
                run = p.add_run(part)
                set_run_font(run, '宋体', 'Times New Roman', 12, bold=False)
        return p

    def add_reference_paragraph(text):
        """Add a reference entry with hanging indent and smaller font."""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        set_hanging_indent_chars(p, 2)
        run = p.add_run(text)
        set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=False)
        return p

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < cols:
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text).strip()
                    cell = table.cell(r_idx, c_idx)
                    cell.text = clean
                    # Set cell font to 宋体五号(10.5pt)
                    for p in cell.paragraphs:
                        p.paragraph_format.line_spacing = 1.0
                        for run in p.runs:
                            set_run_font(run, '宋体', 'Times New Roman', 10.5)
        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip YAML frontmatter
        if i == 0 and line.strip() == '---':
            i += 1
            while i < len(lines) and lines[i].rstrip('\n').strip() != '---':
                i += 1
            i += 1
            continue

        # H1 — 论文标题
        if line.startswith('# ') and not line.startswith('## '):
            flush_table()
            title_text = line[2:].strip()
            p = doc.add_heading('', level=1)
            run = p.add_run(title_text)
            set_run_font(run, '黑体', 'Times New Roman', 22, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # H2 — 一级标题 (## 一、)
        elif line.startswith('## ') and not line.startswith('### '):
            flush_table()
            heading_text = line[3:].strip()
            if '参考文献' in heading_text:
                in_references = True
            p = doc.add_heading('', level=2)
            run = p.add_run(heading_text)
            set_run_font(run, '黑体', 'Times New Roman', 16, bold=True)

        # H3 — 二级标题 (### （一）)
        elif line.startswith('### ') and not line.startswith('#### '):
            flush_table()
            heading_text = line[4:].strip()
            p = doc.add_heading('', level=3)
            run = p.add_run(heading_text)
            set_run_font(run, '黑体', 'Times New Roman', 14, bold=True)

        # H4 — 三级标题 (#### 1.)
        elif line.startswith('#### '):
            flush_table()
            heading_text = line[5:].strip()
            p = doc.add_heading('', level=4)
            run = p.add_run(heading_text)
            set_run_font(run, '宋体', 'Times New Roman', 12, bold=True)

        # Table rows
        elif line.startswith('|'):
            if '---' in line or '===' in line:
                i += 1
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True

        # Empty line
        elif not line.strip():
            flush_table()

        # Reference entries (lines starting with [n])
        elif in_references and re.match(r'^\[\d+\]', line.strip()):
            flush_table()
            add_reference_paragraph(line.strip())

        # Bold-starting paragraph (摘要：/关键词：)
        elif line.strip().startswith('**') and '**' in line.strip()[2:]:
            flush_table()
            add_body_paragraph(line.strip(), indent=True)

        # Regular paragraph
        else:
            flush_table()
            clean = line.strip()
            if clean:
                add_body_paragraph(clean, indent=True)

        i += 1

    flush_table()
    doc.save(docx_path)
    print(f"  ✓ DOCX created: {docx_path}")


# ============================================================
# HTML EXPORT
# ============================================================

def md_to_html(md_path: str, html_path: str):
    """Convert Markdown to standalone HTML with Chinese academic formatting."""
    import markdown

    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Remove YAML frontmatter
    text = re.sub(r'^---\n.*?\n---\n', '', text, flags=re.DOTALL)

    # Convert markdown to HTML
    html_body = markdown.markdown(
        text,
        extensions=['tables', 'toc', 'nl2br', 'sane_lists']
    )

    # Extract title from first H1
    title_match = re.search(r'<h1>(.+?)</h1>', html_body)
    title = title_match.group(1) if title_match else '论文'

    full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        /* ===== 页面 ===== */
        body {{
            font-family: "SimSun", "STSong", "Songti SC", serif;
            font-size: 12pt;
            line-height: 1.5;
            max-width: 210mm;
            margin: 0 auto;
            padding: 2.54cm 3.17cm;
            color: #000;
            background: #fff;
        }}
        /* ===== 标题 ===== */
        h1 {{
            font-family: "SimHei", "Heiti SC", sans-serif;
            font-size: 22pt;
            font-weight: bold;
            text-align: center;
            margin: 0 0 12pt 0;
            line-height: 1.5;
        }}
        h2 {{
            font-family: "SimHei", "Heiti SC", sans-serif;
            font-size: 16pt;
            font-weight: bold;
            text-align: left;
            margin: 6pt 0 6pt 0;
            line-height: 1.5;
            border: none;
            padding: 0;
        }}
        h3 {{
            font-family: "SimHei", "Heiti SC", sans-serif;
            font-size: 14pt;
            font-weight: bold;
            text-align: left;
            margin: 3pt 0 0 0;
            line-height: 1.5;
        }}
        h4 {{
            font-family: "SimSun", "STSong", serif;
            font-size: 12pt;
            font-weight: bold;
            text-align: left;
            margin: 0;
            line-height: 1.5;
        }}
        /* ===== 正文段落: 首行缩进2字符 ===== */
        p {{
            text-indent: 2em;
            margin: 0;
            padding: 0;
            font-size: 12pt;
            line-height: 1.5;
        }}
        /* ===== 摘要/关键词 (加粗开头) ===== */
        p > strong:first-child {{
            font-family: "SimHei", "Heiti SC", sans-serif;
        }}
        /* ===== 表格 ===== */
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 6pt 0;
            font-size: 10.5pt;
        }}
        th {{
            border-top: 1.5pt solid #000;
            border-bottom: 0.75pt solid #000;
            border-left: none;
            border-right: none;
            padding: 4pt 8pt;
            font-family: "SimHei", sans-serif;
            font-weight: bold;
            text-align: center;
        }}
        td {{
            border: none;
            border-bottom: 0.5pt solid #999;
            padding: 4pt 8pt;
            text-align: center;
        }}
        tr:last-child td {{
            border-bottom: 1.5pt solid #000;
        }}
        /* ===== 参考文献: 悬挂缩进 ===== */
        h2 + ol, h2 + ul {{
            padding-left: 2em;
            text-indent: -2em;
            font-size: 10.5pt;
            line-height: 1.5;
            list-style: none;
        }}
        /* ===== 引用块(对话) ===== */
        blockquote {{
            border-left: 3px solid #ccc;
            margin: 6pt 0;
            padding: 4pt 12pt;
            color: #333;
            background: #f9f9f9;
            font-size: 11pt;
        }}
        blockquote p {{
            text-indent: 0;
        }}
        a {{
            color: #000;
            text-decoration: none;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    print(f"  ✓ HTML created: {html_path}")


# ============================================================
# MAIN
# ============================================================

def main():
    if len(sys.argv) < 3:
        print(f"Usage: {sys.argv[0]} <input.md> <output_dir> [formats]")
        print(f"  formats: all,docx,html (default: all)")
        sys.exit(1)

    input_file = sys.argv[1]
    output_dir = sys.argv[2]
    formats = sys.argv[3] if len(sys.argv) > 3 else 'all'

    base_name = Path(input_file).stem
    os.makedirs(output_dir, exist_ok=True)

    format_list = ['docx', 'html'] if formats == 'all' else formats.split(',')

    print(f"=========================================")
    print(f"Exporting: {base_name}")
    print(f"Output dir: {output_dir}")
    print(f"Formats: {', '.join(format_list)}")
    print(f"=========================================")

    for fmt in format_list:
        if fmt == 'docx':
            docx_path = os.path.join(output_dir, f"{base_name}.docx")
            try:
                md_to_docx(input_file, docx_path)
            except Exception as e:
                print(f"  ✗ DOCX failed: {e}")
                import traceback
                traceback.print_exc()
        elif fmt == 'html':
            html_path = os.path.join(output_dir, f"{base_name}.html")
            try:
                md_to_html(input_file, html_path)
            except Exception as e:
                print(f"  ✗ HTML failed: {e}")
        elif fmt == 'pdf':
            print(f"  → PDF: SKIPPED (requires pandoc + LaTeX)")
        else:
            print(f"  Unknown format: {fmt}")

    print(f"\n=========================================")
    print(f"Export complete!")
    print(f"=========================================")
    for f in os.listdir(output_dir):
        if f.startswith(base_name):
            size = os.path.getsize(os.path.join(output_dir, f))
            print(f"  {f} ({size:,} bytes)")


if __name__ == '__main__':
    main()
